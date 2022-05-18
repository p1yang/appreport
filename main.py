#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#@Author: p1yang
#@File: main.py
#@Time: 2022/05/17 08:36:02
import os
from pydoc import Doc
from telnetlib import DO
from click import style
import requests
import re
import urllib
import time
from docx import Document

#
kuchuan_search_app = "https://android.kuchuan.com/searchappallmarket?kw={}&page=1&count=20&market=0&date={}"
kuchuan_search_appInfo = "https://android.kuchuan.com/totaldownload?packagename={}&date={}"
kuchuan_search_developer = "https://android.kuchuan.com/page/detail/download?package={}&infomarketid=10&site=0#!/sum/{}"
kuchuan_info = "https://android.kuchuan.com/page/detail/download?package={}&infomarketid=1&site=0#!/sum/{}"
app_name = "" #软件名
app_version = "" #版本
app_hash = "" #hash
package_name = "" #包名
totla_downloads = 0 #下载量
developer = "" #开发商名称
leak_type = "" #漏洞类型，需要自行填写
is_open = "" #发布情况
hazard_level = ""#危害等级
addr = "" #所属省份

def getDate():
    date = int( round(time.time() * 1000)) - int(int(time.time() - time.timezone) % 86400)*1000
    return date

# 获取总下载量
def getTotalDownloads():
    url = kuchuan_search_appInfo.format(package_name,getDate())
    head_data = {
        'Accept': 'application/json, text/plain, */*',
        'Cookie': 'sign=Bvm8IuYCoa3V%2BWAkK7B3ZIRk8YmTFp2JsVNjN3aG3KHixvMGDZhODl%2FDtkbzQEA%2FTucc51R8NMeqo7T6BDCM%2FOJVBM8SUT9AZkOU8IW3tYmGgbi%2BsTlaUWYdaaqqdkrr3RjsAk2jlrs9%2FFTeaRkclWbGkUh%2BjdosdCuLnkp9FHE%3D; token=97365f64863924b85895bf08512ae318; _ga=GA1.2.1848213340.1652690100; _gid=GA1.2.1156197270.1652690100; Hm_lpvt_ef5592497ef84b441c6f40f98ada9c4f=1652759163; Hm_lvt_ef5592497ef84b441c6f40f98ada9c4f=1652690097,1652756156; Hm_lpvt_84ee2face82d86b7584b5bbee44dfbe7=1652759163; Hm_lvt_84ee2face82d86b7584b5bbee44dfbe7=1652689992,1652756135; CNZZDATA1257619731=708008061-1652680487-https%253A%252F%252Fios.kuchuan.com%252F%7C1652758660; CNZZDATA1260525360=2013768316-1652680496-https%253A%252F%252Fios.kuchuan.com%252F%7C1652759065; CNZZDATA5103679=cnzz_eid%3D77776780-1652681561-https%253A%252F%252Fios.kuchuan.com%252F%26ntime%3D1652758660; accessId=a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900; pageViewNum=34; qimo_seokeywords_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=%E6%9C%AA%E7%9F%A5; qimo_seosource_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=%E5%85%B6%E4%BB%96%E7%BD%91%E7%AB%99; qimo_xstKeywords_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=; _gat=1; _Coolchuan_com_session=HIIK8m9RDfLaZQKJdf%2BQ7vy3i%2FJvHrzxD3AJMvs9Yhl6icfT7JE3zv3bf4F0S5EKbqCNxdsIZg0pOX7UM6asc2oWe27Z6hQq43lUX6X%2FUpGf4U%2Fb2egQl%2Bd9YnANS%2FvQtZhoGacbYbJG2NBBkqagVS4NGu9fPVx0U4TfNvWuMNg%3D; uid=897366; uname=17550953521; _qddaz=QD.6093k.3lnrnj.l39m3m4q; JSESSIONID=uokvxkxcyddv14109a4bgzswq; JSESSIONID=2kw5z7oy0u851is7sz0qtcvcv; href=https%3A%2F%2Fandroid.kuchuan.com%2Fpage%2Fdetail%2FsearchResult%3Fkw%3D%25E8%25B1%25AB%25E7%2595%2585%25E6%25B4%2597%25E5%2590%25A7%26site%3D0%26marketId%3D0; UM_distinctid=180cbfff87121d-0f40ca97baa516-3a62684b-13c680-180cbfff8724ed; uniqueId=daa491ba33f54d68b40d1390b9d57cef',
        'Accept-Language': 'zh-CN,zh-Hans;q=0.9',
        'Host': 'android.kuchuan.com',
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.4 Safari/605.1.15',
        'Referer': 'https://android.kuchuan.com/page/detail/download?package={}&infomarketid=10&site=0'.format(package_name),
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'X-Requested-With': 'XMLHttpRequest'
    }
    data = requests.get(url,headers=head_data).json()
    #data = {"status":200, "msg":"请求成功", "data":{"联想":4,"OPPO":195,"华为":10000,"应用宝":23,"魅族":20,"vivo":1908}}
    nums = data["data"]
    downloads = 0
    for num in nums:
        downloads += nums[num]
    return downloads


#获取软件包名
def getPackageName():
    app_name_hex = urllib.parse.quote(app_name)
    head_data = {
        "Accept": "application/json, text/plain, */*",
        "Cookie": "sign=Bvm8IuYCoa3V%2BWAkK7B3ZIRk8YmTFp2JsVNjN3aG3KHixvMGDZhODl%2FDtkbzQEA%2FTucc51R8NMeqo7T6BDCM%2FOJVBM8SUT9AZkOU8IW3tYmGgbi%2BsTlaUWYdaaqqdkrr3RjsAk2jlrs9%2FFTeaRkclWbGkUh%2BjdosdCuLnkp9FHE%3D; token=97365f64863924b85895bf08512ae318; _ga=GA1.2.1848213340.1652690100; _gid=GA1.2.1156197270.1652690100; Hm_lpvt_ef5592497ef84b441c6f40f98ada9c4f=1652759163; Hm_lvt_ef5592497ef84b441c6f40f98ada9c4f=1652690097,1652756156; Hm_lpvt_84ee2face82d86b7584b5bbee44dfbe7=1652759163; Hm_lvt_84ee2face82d86b7584b5bbee44dfbe7=1652689992,1652756135; CNZZDATA1257619731=708008061-1652680487-https%253A%252F%252Fios.kuchuan.com%252F%7C1652758660; CNZZDATA1260525360=2013768316-1652680496-https%253A%252F%252Fios.kuchuan.com%252F%7C1652759065; CNZZDATA5103679=cnzz_eid%3D77776780-1652681561-https%253A%252F%252Fios.kuchuan.com%252F%26ntime%3D1652758660; accessId=a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900; pageViewNum=34; qimo_seokeywords_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=%E6%9C%AA%E7%9F%A5; qimo_seosource_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=%E5%85%B6%E4%BB%96%E7%BD%91%E7%AB%99; qimo_xstKeywords_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=; _gat=1; _Coolchuan_com_session=HIIK8m9RDfLaZQKJdf%2BQ7vy3i%2FJvHrzxD3AJMvs9Yhl6icfT7JE3zv3bf4F0S5EKbqCNxdsIZg0pOX7UM6asc2oWe27Z6hQq43lUX6X%2FUpGf4U%2Fb2egQl%2Bd9YnANS%2FvQtZhoGacbYbJG2NBBkqagVS4NGu9fPVx0U4TfNvWuMNg%3D; uid=897366; uname=17550953521; _qddaz=QD.6093k.3lnrnj.l39m3m4q; JSESSIONID=uokvxkxcyddv14109a4bgzswq; JSESSIONID=2kw5z7oy0u851is7sz0qtcvcv; href=https%3A%2F%2Fandroid.kuchuan.com%2Fpage%2Fdetail%2FsearchResult%3Fkw%3D%25E8%25B1%25AB%25E7%2595%2585%25E6%25B4%2597%25E5%2590%25A7%26site%3D0%26marketId%3D0; UM_distinctid=180cbfff87121d-0f40ca97baa516-3a62684b-13c680-180cbfff8724ed; uniqueId=daa491ba33f54d68b40d1390b9d57cef",
        "Accept-Language": "zh-CN,zh-Hans;q=0.9",
        "Host": "android.kuchuan.com",
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.4 Safari/605.1.15",
        "Referer": "https://android.kuchuan.com/page/detail/searchResult?kw={}&site=0&marketId=0".format(app_name_hex),
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "X-Requested-With": "XMLHttpRequest"
    }
    url = kuchuan_search_app.format(app_name_hex,getDate())
    data = requests.get(url,headers=head_data).json()
    #data = {'status': 200, 'msg': '请求成功', 'data': {'result': [{'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttp%3A%2F%2Fpp.myapp.com%2Fma_icon%2F0%2Ficon_54127626_1631521246%2F96', 'market_id': '10', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttp%3A%2F%2Fpp.myapp.com%2Fma_icon%2F0%2Ficon_54127626_1631521246%2F96', 'market_id': '10', 'packageName': 'com.benben.zhuangxiugong', 'title': '豫上装'}, 'packageName': 'com.benben.zhuangxiugong', 'title': '豫上装'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F675a7b320a9946248d298b1dfacceddf.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F675a7b320a9946248d298b1dfacceddf.png', 'market_id': '18', 'packageName': 'cn.org.camib.www2.WeBrowser.zszx', 'title': '掌上装修'}, 'packageName': 'cn.org.camib.www2.WeBrowser.zszx', 'title': '掌上装修'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F7281789391134bbaa3a785e0aea27286.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F7281789391134bbaa3a785e0aea27286.png', 'market_id': '18', 'packageName': 'com.cxzg.m.zhuangxzs', 'title': '掌上装修装饰'}, 'packageName': 'com.cxzg.m.zhuangxzs', 'title': '掌上装修装饰'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F6f436ded19bd465ab70130f47b771f1f.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F6f436ded19bd465ab70130f47b771f1f.png', 'market_id': '18', 'packageName': 'com.eggpain.zhangshangzhuangxiushangcheng690', 'title': '掌上装修商城'}, 'packageName': 'com.eggpain.zhangshangzhuangxiushangcheng690', 'title': '掌上装修商城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd0a162e7caf5409eb5bf4f60d235eb7d.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd0a162e7caf5409eb5bf4f60d235eb7d.png', 'market_id': '18', 'packageName': 'cn.org.camib.www2.WeBrowser.zszssc', 'title': '掌上装饰商城'}, 'packageName': 'cn.org.camib.www2.WeBrowser.zszssc', 'title': '掌上装饰商城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd85271348e034d89919d36764470b07e_1.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd85271348e034d89919d36764470b07e_1.png', 'market_id': '18', 'packageName': 'com.ddm.shoppingmall', 'title': '掌上装饰城'}, 'packageName': 'com.ddm.shoppingmall', 'title': '掌上装饰城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F7e403bf864314815a62bd50be4aa10de.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F7e403bf864314815a62bd50be4aa10de.png', 'market_id': '18', 'packageName': 'cn.org.camib.www2.WeBrowser.zszxwfns', 'title': '掌上装修网'}, 'packageName': 'cn.org.camib.www2.WeBrowser.zszxwfns', 'title': '掌上装修网'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttp%3A%2F%2Fimg.wdjimg.com%2Fmms%2Ficon%2Fv1%2Ff%2Fcb%2F77f1191ff7b74553597ceacd435aacbf_256_256.png', 'market_id': '17', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttp%3A%2F%2Fimg.wdjimg.com%2Fmms%2Ficon%2Fv1%2Ff%2Fcb%2F77f1191ff7b74553597ceacd435aacbf_256_256.png', 'market_id': '17', 'packageName': 'com.sino.app.advancedB31317', 'title': '掌上装饰城'}, 'packageName': 'com.sino.app.advancedB31317', 'title': '掌上装饰城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd75fde337e2541e5b02a29402ab9fd67.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd75fde337e2541e5b02a29402ab9fd67.png', 'market_id': '18', 'packageName': 'cn.org.camib.WeBrowser.zszsc', 'title': '掌上装饰城'}, 'packageName': 'cn.org.camib.WeBrowser.zszsc', 'title': '掌上装饰城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F686b52b6d73d40ce86393ef91a255c06.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F686b52b6d73d40ce86393ef91a255c06.png', 'market_id': '18', 'packageName': 'cn.org.camib.www2.WeBrowser.zszsw8', 'title': '掌上装饰网'}, 'packageName': 'cn.org.camib.www2.WeBrowser.zszsw8', 'title': '掌上装饰网'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F2ac08f23a57c453f9a052e6ee202a6b5.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F2ac08f23a57c453f9a052e6ee202a6b5.png', 'market_id': '18', 'packageName': 'com.zdweilai.WeBrowser.zszslm', 'title': '掌上装饰联盟'}, 'packageName': 'com.zdweilai.WeBrowser.zszslm', 'title': '掌上装饰联盟'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F5603d6cc7a3e4e4e81e824c486fd64d3.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F5603d6cc7a3e4e4e81e824c486fd64d3.png', 'market_id': '18', 'packageName': 'com.eggpain.zhangshangzhuangxiumenhu3911', 'title': '掌上装修门户'}, 'packageName': 'com.eggpain.zhangshangzhuangxiumenhu3911', 'title': '掌上装修门户'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fecee3ec62ec6442d9038b9fb7cbe1fda.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fecee3ec62ec6442d9038b9fb7cbe1fda.png', 'market_id': '18', 'packageName': 'com.eggpain.zhangshangzhuangshigongcheng2508', 'title': '掌上装饰工程'}, 'packageName': 'com.eggpain.zhangshangzhuangshigongcheng2508', 'title': '掌上装饰工程'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fc34f769e8be64097b002f5bc76ac0317.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fc34f769e8be64097b002f5bc76ac0317.png', 'market_id': '18', 'packageName': 'com.yes.bokai', 'title': '博凯掌上装饰'}, 'packageName': 'com.yes.bokai', 'title': '博凯掌上装饰'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd3f85eda9f344d119c78f116de95b7e2.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fd3f85eda9f344d119c78f116de95b7e2.png', 'market_id': '18', 'packageName': 'com.eggpain.zhangshangzhuangxiudaquan1510', 'title': '掌上装修大全'}, 'packageName': 'com.eggpain.zhangshangzhuangxiudaquan1510', 'title': '掌上装修大全'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F556f81a1bfcd48a28422d11dc3053bbe.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F556f81a1bfcd48a28422d11dc3053bbe.png', 'market_id': '18', 'packageName': 'com.eggpain.zhangshangzhuangshicailiaocheng3065', 'title': '掌上装饰材料城'}, 'packageName': 'com.eggpain.zhangshangzhuangshicailiaocheng3065', 'title': '掌上装饰材料城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F51630f263a0546ca9156b053ec506aaa.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2F51630f263a0546ca9156b053ec506aaa.png', 'market_id': '18', 'packageName': 'com.eggpain.zhangshangzhuangshicailiaowang570', 'title': '掌上装饰材料网'}, 'packageName': 'com.eggpain.zhangshangzhuangshicailiaowang570', 'title': '掌上装饰材料网'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fc144aebbaece4e31908a15d19860e4e9.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fc144aebbaece4e31908a15d19860e4e9.png', 'market_id': '18', 'packageName': 'com.wise.zhshzhuangshgchwang', 'title': '掌上装饰工程网'}, 'packageName': 'com.wise.zhshzhuangshgchwang', 'title': '掌上装饰工程网'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fb96919e0c9b144ffbb3d514005ec60d7.png', 'market_id': '18', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttps%3A%2F%2Fappimg.dbankcdn.com%2Fapplication%2Ficon144%2Fb96919e0c9b144ffbb3d514005ec60d7.png', 'market_id': '18', 'packageName': 'com.zhongsou.zmall.zszsclscmall', 'title': '掌上装饰材料商城'}, 'packageName': 'com.zhongsou.zmall.zszsclscmall', 'title': '掌上装饰材料商城'}, {'baikeActive': 0, 'baikeUrl': '', 'IsAuth': 0, 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttp%3A%2F%2Fp2.qhimg.com%2Ft0112fc94203ebfc670.png', 'market_id': '1', 'appinfo': {'baikeActive': '0', 'baikeUrl': '', 'IsAuth': '0', 'icon': 'https://cdn-image.kuchuan.com/transferPicUrl?picUrl=http%3A%2F%2Fcdn-image.kuchuan.com%2FtransferPicUrl%3FpicUrl%3Dhttp%3A%2F%2Fp2.qhimg.com%2Ft0112fc94203ebfc670.png', 'market_id': '1', 'packageName': 'com.rd.huatest', 'title': '掌上装裱工-摩奈特装裱'}, 'packageName': 'com.rd.huatest', 'title': '掌上装裱工-摩奈特装裱'}], 'maxCount': 41787, 'isEnd': False}}
    infos = data['data']['result']
    package_name = ""
    for info in infos:
        if app_name == info["title"]:
            package_name = info["packageName"]
            break
    return package_name

#获取开发商名称
def getDeveloper():
    url = kuchuan_search_developer.format(package_name,package_name)
    head_data = {
        'Cookie': 'sign=Bvm8IuYCoa3V%2BWAkK7B3ZIRk8YmTFp2JsVNjN3aG3KHixvMGDZhODl%2FDtkbzQEA%2FTucc51R8NMeqo7T6BDCM%2FOJVBM8SUT9AZkOU8IW3tYmGgbi%2BsTlaUWYdaaqqdkrr3RjsAk2jlrs9%2FFTeaRkclWbGkUh%2BjdosdCuLnkp9FHE%3D; token=97365f64863924b85895bf08512ae318; _ga=GA1.2.1848213340.1652690100; _gid=GA1.2.1156197270.1652690100; Hm_lpvt_ef5592497ef84b441c6f40f98ada9c4f=1652759163; Hm_lvt_ef5592497ef84b441c6f40f98ada9c4f=1652690097,1652756156; Hm_lpvt_84ee2face82d86b7584b5bbee44dfbe7=1652759163; Hm_lvt_84ee2face82d86b7584b5bbee44dfbe7=1652689992,1652756135; CNZZDATA1257619731=708008061-1652680487-https%253A%252F%252Fios.kuchuan.com%252F%7C1652758660; CNZZDATA1260525360=2013768316-1652680496-https%253A%252F%252Fios.kuchuan.com%252F%7C1652759065; CNZZDATA5103679=cnzz_eid%3D77776780-1652681561-https%253A%252F%252Fios.kuchuan.com%252F%26ntime%3D1652758660; accessId=a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900; pageViewNum=34; qimo_seokeywords_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=%E6%9C%AA%E7%9F%A5; qimo_seosource_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=%E5%85%B6%E4%BB%96%E7%BD%91%E7%AB%99; qimo_xstKeywords_a1cb1fb0-0f1e-11e9-9e27-913eb5b7d900=; _gat=1; _Coolchuan_com_session=HIIK8m9RDfLaZQKJdf%2BQ7vy3i%2FJvHrzxD3AJMvs9Yhl6icfT7JE3zv3bf4F0S5EKbqCNxdsIZg0pOX7UM6asc2oWe27Z6hQq43lUX6X%2FUpGf4U%2Fb2egQl%2Bd9YnANS%2FvQtZhoGacbYbJG2NBBkqagVS4NGu9fPVx0U4TfNvWuMNg%3D; uid=897366; uname=17550953521; _qddaz=QD.6093k.3lnrnj.l39m3m4q; JSESSIONID=uokvxkxcyddv14109a4bgzswq; JSESSIONID=2kw5z7oy0u851is7sz0qtcvcv; href=https%3A%2F%2Fandroid.kuchuan.com%2Fpage%2Fdetail%2FsearchResult%3Fkw%3D%25E8%25B1%25AB%25E7%2595%2585%25E6%25B4%2597%25E5%2590%25A7%26site%3D0%26marketId%3D0; UM_distinctid=180cbfff87121d-0f40ca97baa516-3a62684b-13c680-180cbfff8724ed; uniqueId=daa491ba33f54d68b40d1390b9d57cef',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Host': 'android.kuchuan.com',
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.4 Safari/605.1.15',
        'Accept-Language': 'zh-CN,zh-Hans;q=0.9',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive'
    }
    data = requests.get(url,headers=head_data).text
    s = re.findall(r'"developer":"[\u4e00-\u9fa5]+"',data)
    return s[0].split('"')[-2]

#获取hash，使用的是md5
def getHash(fileName):
    res = os.popen("md5 "+fileName).read()
    return res[res.index("=")+2:-1]
    
def printInfo():
    print("[软件名]："+app_name)        
    print("[软件版本]："+app_version)
    print("[软件哈希]："+app_hash)
    print("[软件包名]："+package_name)
    print("[下载量]："+str(totla_downloads))
    print("[开发商]："+developer)

def createReport():
    newDoc = Document()
    newDoc.add_heading("{} {}APP {}版本存在{}问题".format(developer,app_name,app_version,leak_type),level=0)
    newDoc.add_paragraph("一、受影响产品名称：{}{}APP {} 安卓版".format(developer,app_name,app_version))
    newDoc.add_paragraph("二、网络产品提供者/影响企业的所属省份：{}".format(addr))
    newDoc.add_paragraph("三、APP包名：{}".format(package_name))
    newDoc.add_paragraph("四、APP哈希值：MD5:{}".format(app_hash))
    newDoc.add_paragraph("五、漏洞成因类型：{}问题".format(leak_type))
    newDoc.add_paragraph("六、发布情况：{}".format(is_open))
    newDoc.add_paragraph("七、危害等级：{}".format(hazard_level))
    newDoc.add_paragraph("八、总下载量：{}".format(str(totla_downloads)))
    newDoc.add_paragraph("这里放截图")
    newDoc.add_paragraph("九、漏洞简介\n")
    newDoc.save("{}/{}存在{}问题.docx".format(app_name,app_name,leak_type))

if __name__ == "__main__":
    print("[*]开始获取app信息")
    s = os.listdir()
    name = ""
    for i in s:
        if ".apk" in i:
            name = i
            app_name = i[:i.index("_")]
            app_version = i[i.index("_")+1:i.index("apk")-1]
            app_hash = getHash(i)
            break
    package_name = getPackageName()
    if package_name != "":
        totla_downloads = getTotalDownloads()
        developer = getDeveloper()
    else:
        print("未查到该软件，请手动查询")
    printInfo()
    print("下面请输入其他信息")
    addr = input("[所属省份]：")
    leak_type = input("[漏洞类型]：")
    is_open = input("[发布情况]：")
    hazard_level = input("[危害等级]：")
    print("[*]开始生成报告")
    os.mkdir(app_name)
    createReport()
    os.system("mv {} ./{}/".format(i,app_name))
    print("[*]报告生成完毕，手动补充漏洞截图和简介")
    print(kuchuan_info.format(package_name,package_name))